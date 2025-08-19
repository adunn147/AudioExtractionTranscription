import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import subprocess
import threading
import json
import time
import datetime
from pathlib import Path
import sys

# Try to import optional dependencies
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False
    print("Warning: tkinterdnd2 not available. Drag & drop will be disabled.")

try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False
    print("Warning: NumPy not available. This may cause issues with audio processing.")

try:
    import whisper
    HAS_WHISPER = True
except ImportError:
    HAS_WHISPER = False
    print("Warning: OpenAI Whisper not available. Transcription will be disabled.")

# WhisperX removed - using standard Whisper only

class VideoTranscriberApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Video Audio Extractor & Transcriber")
        # Fixed: Increased window size for better visibility
        self.root.geometry("750x850")
        
        # Variables
        self.video_file = tk.StringVar()
        self.output_folder = tk.StringVar()
        self.extract_audio = tk.BooleanVar(value=True)
        self.audio_format = tk.StringVar(value="mp3")
        self.transcribe_audio = tk.BooleanVar(value=False)
        self.whisper_model = tk.StringVar(value="base")
        self.number_formatting = tk.BooleanVar(value=False)
        self.processing = False
        
        # Whisper model info
        self.model_info = {
            "tiny": ("Fastest, least accurate (~2-3x realtime)", "~20-30 min for 1 hour"),
            "base": ("Good balance of speed/accuracy (~1-2x realtime)", "~30-60 min for 1 hour"),
            "small": ("Better accuracy, slower (~0.5-1x realtime)", "~60-120 min for 1 hour"),
            "medium": ("High accuracy, much slower (~0.2-0.5x realtime)", "~2-5 hours for 1 hour"),
            "large": ("Highest accuracy, very slow (~0.1-0.3x realtime)", "~3-10 hours for 1 hour")
        }
        
        self.setup_ui()
        self.check_startup_dependencies()
        
    def check_startup_dependencies(self):
        """Check dependencies at startup and show warnings"""
        warnings = []
        
        if not HAS_NUMPY:
            warnings.append("NumPy is missing. Install with: pip install --user numpy")
        
        if not HAS_WHISPER:
            warnings.append("OpenAI Whisper is missing. Install with: pip install --user openai-whisper")
            # Disable transcription options
            self.transcribe_audio.set(False)
            
        if not HAS_DND:
            warnings.append("tkinterdnd2 is missing. Drag & drop will be unavailable.")
            
        if warnings:
            warning_msg = "Missing Dependencies:\n\n" + "\n".join(f"• {w}" for w in warnings)
            warning_msg += "\n\nRun 'python setup_dependencies.py' to install missing packages."
            messagebox.showwarning("Dependency Warning", warning_msg)
        
    def setup_ui(self):
        # Main frame with scrollable canvas for better layout
        canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Main frame
        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        row = 0
        
        # Video file selection
        ttk.Label(main_frame, text="Video File:").grid(row=row, column=0, sticky=tk.W, pady=5)
        
        video_frame = ttk.Frame(main_frame)
        video_frame.grid(row=row, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        video_frame.columnconfigure(0, weight=1)
        
        self.video_entry = ttk.Entry(video_frame, textvariable=self.video_file, width=50)
        self.video_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(video_frame, text="Browse", command=self.browse_video).grid(row=0, column=1)
        
        # Drag and drop area (only if DND is available)
        if HAS_DND:
            row += 1
            drop_frame = ttk.LabelFrame(main_frame, text="Drag & Drop Video Here", padding="20")
            drop_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
            drop_frame.columnconfigure(0, weight=1)
            
            self.drop_label = ttk.Label(drop_frame, text="Drop video file here or use Browse button above", 
                                       anchor="center", background="lightgray")
            self.drop_label.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=20)
            
            # Enable drag and drop
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self.on_drop)
        else:
            # Show message that drag & drop is unavailable
            row += 1
            info_frame = ttk.Frame(main_frame)
            info_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
            
            ttk.Label(info_frame, text="Drag & Drop unavailable (tkinterdnd2 not installed)", 
                     foreground="orange").grid(row=0, column=0)
        
        # Output folder
        row += 1
        ttk.Label(main_frame, text="Output Folder:").grid(row=row, column=0, sticky=tk.W, pady=5)
        
        output_frame = ttk.Frame(main_frame)
        output_frame.grid(row=row, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        output_frame.columnconfigure(0, weight=1)
        
        self.output_entry = ttk.Entry(output_frame, textvariable=self.output_folder, width=50)
        self.output_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(output_frame, text="Browse", command=self.browse_output).grid(row=0, column=1)
        
        # Audio extraction options
        row += 1
        audio_frame = ttk.LabelFrame(main_frame, text="Audio Extraction", padding="10")
        audio_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        ttk.Checkbutton(audio_frame, text="Extract Audio", variable=self.extract_audio,
                       command=self.on_extract_audio_change).grid(row=0, column=0, sticky=tk.W)
        
        ttk.Label(audio_frame, text="Format:").grid(row=0, column=1, padx=(20, 5))
        self.audio_format_combo = ttk.Combobox(audio_frame, textvariable=self.audio_format, 
                                        values=["mp3", "wav"], state="readonly", width=10)
        self.audio_format_combo.grid(row=0, column=2)
        
        # Transcription options
        row += 1
        transcribe_frame = ttk.LabelFrame(main_frame, text="Audio Transcription", padding="10")
        transcribe_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        transcribe_frame.columnconfigure(1, weight=1)
        
        self.transcribe_checkbox = ttk.Checkbutton(transcribe_frame, text="Transcribe Audio", 
                                                   variable=self.transcribe_audio,
                                                   command=self.on_transcribe_change)
        self.transcribe_checkbox.grid(row=0, column=0, columnspan=3, sticky=tk.W)
        
        if not HAS_WHISPER:
            self.transcribe_checkbox.configure(state="disabled")
            ttk.Label(transcribe_frame, text="(Requires: pip install --user openai-whisper)", 
                     foreground="red", font=("TkDefaultFont", 8)).grid(row=0, column=3, padx=10)
        
        # Whisper model selection
        ttk.Label(transcribe_frame, text="Whisper Model:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.model_combo = ttk.Combobox(transcribe_frame, textvariable=self.whisper_model,
                                       values=list(self.model_info.keys()), state="disabled", width=15)
        self.model_combo.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        self.model_combo.bind('<<ComboboxSelected>>', self.update_model_info)
        
        # Model info display
        self.model_info_label = ttk.Label(transcribe_frame, text=self.model_info["base"][0], 
                                         foreground="blue", wraplength=400)
        self.model_info_label.grid(row=1, column=2, sticky=tk.W, padx=10, pady=5)
        
        self.model_time_label = ttk.Label(transcribe_frame, text=f"Est. time: {self.model_info['base'][1]}", 
                                         foreground="darkgreen")
        self.model_time_label.grid(row=2, column=1, columnspan=2, sticky=tk.W, padx=5)
        
        # Number formatting option (separate section with proper spacing)
        formatting_frame = ttk.Frame(transcribe_frame)
        formatting_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(15, 5))
        
        self.number_checkbox = ttk.Checkbutton(formatting_frame, text="Number formatting (e.g., 'five percent' → '5%')", 
                                               variable=self.number_formatting, state="disabled")
        self.number_checkbox.grid(row=0, column=0, sticky=tk.W)
        
        # Progress and status
        row += 1
        status_frame = ttk.LabelFrame(main_frame, text="Status", padding="10")
        status_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        status_frame.columnconfigure(0, weight=1)
        
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate')
        self.progress.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=5)
        
        self.status_label = ttk.Label(status_frame, text="Ready")
        self.status_label.grid(row=1, column=0, sticky=tk.W)
        
        # Log area
        log_frame = ttk.LabelFrame(status_frame, text="Log", padding="5")
        log_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=5)
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        
        self.log_text = tk.Text(log_frame, height=8, wrap=tk.WORD)
        log_scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_scrollbar.set)
        
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Buttons
        row += 1
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=row, column=0, columnspan=3, pady=20)
        
        self.process_button = ttk.Button(button_frame, text="Start Processing", 
                                        command=self.start_processing, style="Accent.TButton")
        self.process_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="Clear Log", command=self.clear_log).pack(side=tk.LEFT, padx=5)
        
        # Pack the canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Initialize states
        self.on_extract_audio_change()
        self.on_transcribe_change()
        
    def on_drop(self, event):
        if HAS_DND:
            files = self.root.tk.splitlist(event.data)
            if files:
                self.video_file.set(files[0])
                self.set_default_output_folder()
            
    def browse_video(self):
        file_path = filedialog.askopenfilename(
            title="Select Video File",
            filetypes=[
                ("Video files", "*.mp4 *.avi *.mov *.mkv *.wmv *.flv *.webm *.m4v"),
                ("All files", "*.*")
            ]
        )
        if file_path:
            self.video_file.set(file_path)
            self.set_default_output_folder()
            
    def browse_output(self):
        folder_path = filedialog.askdirectory(title="Select Output Folder")
        if folder_path:
            self.output_folder.set(folder_path)
            
    def set_default_output_folder(self):
        if self.video_file.get() and not self.output_folder.get():
            video_dir = os.path.dirname(self.video_file.get())
            self.output_folder.set(video_dir)
            
    def on_extract_audio_change(self):
        state = "normal" if self.extract_audio.get() else "disabled"
        self.audio_format_combo.configure(state=state)
        
    def on_transcribe_change(self):
        """Enable/disable transcription-related controls"""
        if HAS_WHISPER and self.transcribe_audio.get():
            # Enable transcription controls
            self.model_combo.configure(state="readonly")
            self.number_checkbox.configure(state="normal")
        else:
            # Disable transcription controls
            self.model_combo.configure(state="disabled")
            self.number_checkbox.configure(state="disabled")
        
    def update_model_info(self, event=None):
        model = self.whisper_model.get()
        if model in self.model_info:
            info, time_est = self.model_info[model]
            self.model_info_label.configure(text=info)
            self.model_time_label.configure(text=f"Est. time: {time_est}")
            
    def log_message(self, message):
        self.log_text.insert(tk.END, f"{time.strftime('%H:%M:%S')} - {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        
    def clear_log(self):
        self.log_text.delete(1.0, tk.END)
        
    def update_status(self, message):
        self.status_label.configure(text=message)
        self.log_message(message)
        
    def validate_inputs(self):
        if not self.video_file.get():
            messagebox.showerror("Error", "Please select a video file.")
            return False
            
        if not os.path.exists(self.video_file.get()):
            messagebox.showerror("Error", "Video file does not exist.")
            return False
            
        if not self.output_folder.get():
            messagebox.showerror("Error", "Please select an output folder.")
            return False
            
        if not self.extract_audio.get() and not self.transcribe_audio.get():
            messagebox.showerror("Error", "Please select at least one operation: audio extraction or transcription.")
            return False
            
        return True
        
    def check_dependencies(self):
        """Check if required dependencies are available"""
        missing = []
        
        # Check for ffmpeg
        ffmpeg_path = self.find_ffmpeg()
        if not ffmpeg_path:
            missing.append("ffmpeg.exe (download from https://ffmpeg.org and place in script directory)")
            
        # Check for Python packages
        if not HAS_NUMPY:
            missing.append("numpy (pip install --user numpy)")
            
        if not HAS_WHISPER and self.transcribe_audio.get():
            missing.append("openai-whisper (pip install --user openai-whisper)")
                
        return missing
        
    def find_ffmpeg(self):
        """Find ffmpeg executable"""
        # Check current directory first
        current_dir = os.path.dirname(os.path.abspath(__file__))
        ffmpeg_local = os.path.join(current_dir, "ffmpeg.exe")
        if os.path.exists(ffmpeg_local):
            return ffmpeg_local
            
        # Check if ffmpeg is in PATH
        try:
            result = subprocess.run(["ffmpeg", "-version"], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                return "ffmpeg"
        except:
            pass
            
        return None
        
    def start_processing(self):
        if self.processing:
            messagebox.showwarning("Warning", "Processing is already in progress.")
            return
            
        if not self.validate_inputs():
            return
            
        # Check dependencies
        missing_deps = self.check_dependencies()
        if missing_deps:
            messagebox.showerror("Missing Dependencies", 
                               f"Missing required dependencies:\n\n" + "\n".join(missing_deps))
            return
            
        # Start processing in a separate thread
        self.processing = True
        self.process_button.configure(text="Processing...", state="disabled")
        self.progress.start()
        
        thread = threading.Thread(target=self.process_video)
        thread.daemon = True
        thread.start()
        
    def process_video(self):
        try:
            video_path = self.video_file.get()
            output_dir = self.output_folder.get()
            base_name = os.path.splitext(os.path.basename(video_path))[0]
            
            audio_file = None
            
            # Extract audio if requested
            if self.extract_audio.get():
                self.update_status("Extracting audio...")
                audio_file = self.extract_audio_from_video(video_path, output_dir, base_name)
                
            # Transcribe if requested
            if self.transcribe_audio.get():
                if not audio_file:
                    # Need to extract audio for transcription
                    self.update_status("Extracting audio for transcription...")
                    temp_audio = os.path.join(output_dir, f"{base_name}_temp.wav")
                    self.extract_audio_from_video(video_path, output_dir, f"{base_name}_temp", "wav")
                    audio_file = temp_audio
                    
                self.update_status("Starting transcription...")
                self.transcribe_audio_file(audio_file, output_dir, base_name)
                
                # Clean up temporary file if created
                temp_audio = os.path.join(output_dir, f"{base_name}_temp.wav")
                if os.path.exists(temp_audio):
                    try:
                        os.remove(temp_audio)
                        self.log_message("Cleaned up temporary audio file")
                    except:
                        pass  # Ignore cleanup errors
                    
            self.update_status("Processing completed successfully!")
            messagebox.showinfo("Success", "Processing completed successfully!")
            
        except Exception as e:
            error_msg = f"Error during processing: {str(e)}"
            self.update_status(error_msg)
            self.log_message(f"Full error details: {type(e).__name__}: {e}")
            messagebox.showerror("Error", error_msg)
            
        finally:
            self.processing = False
            self.progress.stop()
            self.process_button.configure(text="Start Processing", state="normal")
            
    def extract_audio_from_video(self, video_path, output_dir, base_name, format_override=None):
        """Extract audio from video using ffmpeg"""
        audio_format = format_override or self.audio_format.get()
        audio_file = os.path.join(output_dir, f"{base_name}.{audio_format}")
        
        ffmpeg_path = self.find_ffmpeg()
        
        cmd = [
            ffmpeg_path,
            "-i", video_path,
            "-vn",  # No video
            "-acodec", "libmp3lame" if audio_format == "mp3" else "pcm_s16le",
            "-ar", "16000",  # 16kHz sample rate (good for Whisper)
            "-ac", "1",  # Mono
            "-y",  # Overwrite output file
            audio_file
        ]
        
        self.log_message(f"Running: {' '.join(cmd)}")
        
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            raise Exception(f"FFmpeg error: {stderr}")
            
        self.log_message(f"Audio extracted: {audio_file}")
        return audio_file
        
    def transcribe_audio_file(self, audio_file, output_dir, base_name):
        """Transcribe audio using Whisper"""
        if not HAS_WHISPER:
            raise Exception("OpenAI Whisper is not installed. Install with: pip install --user openai-whisper")
            
        import whisper
        
        model_name = self.whisper_model.get()
        self.transcribe_with_whisper(audio_file, output_dir, base_name, model_name)
            
    def transcribe_with_whisper(self, audio_file, output_dir, base_name, model_name):
        """Transcribe using standard Whisper"""
        import whisper
        
        self.update_status(f"Loading Whisper model: {model_name}")
        model = whisper.load_model(model_name)
        
        self.update_status("Transcribing audio...")
        result = model.transcribe(audio_file)
        
        # Always create Word document with both clean text and timestamps
        docx_file = os.path.join(output_dir, f"{base_name}_transcript.docx")
        self.create_docx_transcript(result, docx_file, base_name)
        self.log_message(f"Word document saved: {docx_file}")
            
    def format_timestamp(self, seconds):
        """Format seconds as MM:SS or HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        else:
            return f"{minutes:02d}:{seconds:02d}"
            
    def create_docx_transcript(self, result, docx_file, base_name):
        """Create a formatted Word document with both clean text and timestamped transcript"""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            import datetime
        except ImportError:
            self.log_message("Warning: python-docx not installed. Cannot create Word document.")
            self.log_message("Install with: pip install --user python-docx")
            return
            
        # Create document
        doc = Document()
        
        # Set document properties with current date/time
        core_props = doc.core_properties
        core_props.title = f'Transcript: {base_name}'
        core_props.author = 'Video Transcriber'
        core_props.created = datetime.datetime.now()
        core_props.modified = datetime.datetime.now()
        
        # Add title
        title = doc.add_heading(f'Transcript: {base_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        if 'duration' in result:
            doc.add_paragraph(f"Duration: {self.format_timestamp(result['duration'])}")
        doc.add_paragraph(f"Language: {result.get('language', 'Unknown')}")
        doc.add_paragraph("")  # Empty line
        
        # Add clean text section
        clean_heading = doc.add_heading('Clean Text', level=1)
        clean_para = doc.add_paragraph(result["text"])
        
        # Add page break and timestamped section
        doc.add_page_break()
        timestamp_heading = doc.add_heading('Timestamped Transcript', level=1)
        
        # Create a table for better timestamp alignment
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.2) if column == table.columns[0] else Inches(5.5)
        
        for segment in result["segments"]:
            start_time = self.format_timestamp(segment["start"])
            end_time = self.format_timestamp(segment["end"])
            text = segment["text"].strip()
            
            # Add row to table
            row = table.add_row()
            
            # Timestamp cell (left column)
            timestamp_cell = row.cells[0]
            timestamp_para = timestamp_cell.paragraphs[0]
            timestamp_run = timestamp_para.add_run(f"{start_time}\n{end_time}")
            timestamp_run.bold = True
            timestamp_run.font.size = Pt(9)
            try:
                timestamp_run.font.color.rgb = RGBColor(0, 0, 150)  # Dark blue
            except:
                pass  # Color might not work on all systems
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Text cell (right column)
            text_cell = row.cells[1]
            text_para = text_cell.paragraphs[0]
            text_para.add_run(text)
            text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add some padding to cells
            for cell in row.cells:
                cell.vertical_alignment = 1  # Top alignment
        
        # Save document
        doc.save(docx_file)
    
    # Fixed: Removed the orphaned SRT code that was causing the NameError
    # The problematic code was after doc.save(docx_file) and referenced undefined 'srt_file'


def main():
    # Create root window with or without DND support
    if HAS_DND:
        try:
            root = TkinterDnD.Tk()
        except:
            print("Warning: TkinterDnD initialization failed, falling back to regular tkinter")
            root = tk.Tk()
    else:
        root = tk.Tk()
        
    app = VideoTranscriberApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()